"""
Document processing service — handles PDF, EPUB, DOCX, TXT extraction.
Runs as a Celery background task after upload.

Reading order strategy:
- PDF: OCR is the PRIMARY extraction method for every page using Tesseract
  with psm=3 (auto layout detection). This captures ALL visible text including
  text in boxes, tables, sidebars, callouts, and decorative elements — exactly
  what a human eye sees, in the order a human would read it.
  PyMuPDF is used only to extract embedded images (not text).
- EPUB/DOCX/TXT: natural document order preserved via HTML block tag parsing.
"""

import re
import io
import base64
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF — used for image extraction and page rendering only
from PIL import Image
import pytesseract

from core.celery_app import celery_app
from core.config import settings


# ── Main entry point ──────────────────────────────────────────────────────────

def process_document(file_bytes: bytes, file_type: str) -> dict:
    """Route to the appropriate parser based on file type."""
    if file_type == "pdf":
        return _process_pdf(file_bytes)
    elif file_type in ("epub", "application/epub+zip"):
        return _process_epub(file_bytes)
    elif file_type in ("docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
        return _process_docx(file_bytes)
    elif file_type in ("txt", "text/plain"):
        return _process_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


# ── PDF ───────────────────────────────────────────────────────────────────────

def _process_pdf(file_bytes: bytes) -> dict:
    """
    Process PDF using OCR as the primary text extraction method.

    Every page is rendered to an image at 300 DPI and passed through Tesseract
    with psm=3 (automatic page segmentation). This ensures:
    - Text in boxes, tables, sidebars, callouts is captured
    - Reading order matches what a human sees (Tesseract handles layout)
    - Scanned and digital PDFs are treated identically
    - No text is missed due to PDF encoding quirks

    Images are extracted separately via PyMuPDF's image extraction API.
    """
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages = []
    total_words = 0

    for page_num, page in enumerate(doc):
        # ── Primary: Try extracting digital text first ──────────────────────────
        paragraphs = _extract_digital_text(page)
        
        # ── Fallback: If no/little text found, assume scanned and OCR ───────────
        if not paragraphs or sum(p["word_count"] for p in paragraphs) < 15:
            paragraphs = _ocr_page_primary(page)

        total_words += sum(p["word_count"] for p in paragraphs)

        # ── Extract embedded images via PyMuPDF ───────────────────────────────
        images = _extract_page_images(doc, page)

        pages.append({
            "page": page_num + 1,
            "paragraphs": paragraphs,
            "images": images,
        })

    return {
        "pages": pages,
        "total_pages": len(doc),
        "total_words": total_words,
        "toc": _extract_pdf_toc(doc),
    }


def _extract_digital_text(page) -> list:
    """Extract digital text from a PDF page using PyMuPDF blocks."""
    blocks = page.get_text("blocks")
    paragraphs = []
    
    # Sort blocks vertically, then horizontally to maintain reading order
    blocks.sort(key=lambda b: (b[1], b[0]))
    
    for b in blocks:
        if b[6] == 0:  # text block
            text = b[4].strip()
            # Collapse multiple spaces and newlines into single spaces
            text = re.sub(r"\n+", " ", text)
            text = re.sub(r" {2,}", " ", text).strip()
            
            # Skip isolated page numbers / noise
            if len(text) < 3 or re.match(r"^[-–—\s]*\d{1,4}[-–—\s]*$", text):
                continue
                
            paragraphs.append({
                "text": text,
                "bbox": [b[0], b[1], b[2], b[3]],
                "is_heading": _looks_like_heading(text),
                "word_count": len(text.split()),
                "ocr": False,
            })
    return paragraphs


def _ocr_page_primary(page) -> list:
    """
    OCR a PDF page as the primary text extraction method.

    Uses Tesseract psm=3 (fully automatic page segmentation + layout analysis)
    which detects columns, text boxes, tables, and reading order automatically.

    Returns a list of paragraph dicts sorted in reading order.
    """
    # Render page to image at 300 DPI for good OCR accuracy
    pix = page.get_pixmap(dpi=300)
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    pdf_w = page.rect.width   # PDF points
    pdf_h = page.rect.height

    # Try structured extraction first (gives paragraph-level grouping + bboxes)
    paragraphs = _ocr_structured(img, pdf_width=pdf_w, pdf_height=pdf_h)
    if paragraphs:
        return paragraphs

    # Fallback: full-page string split into paragraphs by blank lines
    return _ocr_simple(img)


def _ocr_structured(img: Image.Image, pdf_width: float = 0, pdf_height: float = 0) -> list:
    """
    Use pytesseract.image_to_data to get word-level data with block/paragraph
    grouping. Groups words into paragraphs, sorts by vertical position.
    Filters low-confidence noise (conf < 30).
    Also computes a bbox [x0, y0, x1, y1] in PDF coordinate space so the
    frontend can render highlight overlays on the page image.
    """
    try:
        data = pytesseract.image_to_data(
            img,
            config="--psm 3 --oem 3",
            output_type=pytesseract.Output.DICT,
        )

        img_w, img_h = img.size  # pixel dimensions at 300 DPI

        # Group words by (block_num, par_num) — Tesseract's own paragraph detection
        paragraphs_map: dict = {}
        n = len(data["text"])
        for i in range(n):
            word = data["text"][i].strip()
            conf = int(data["conf"][i])
            if not word or conf < 30:
                continue
            key = (data["block_num"][i], data["par_num"][i])
            x, y, w, h = data["left"][i], data["top"][i], data["width"][i], data["height"][i]
            if key not in paragraphs_map:
                paragraphs_map[key] = {
                    "words": [],
                    "top": y,
                    "x0": x, "y0": y,
                    "x1": x + w, "y1": y + h,
                }
            else:
                paragraphs_map[key]["x0"] = min(paragraphs_map[key]["x0"], x)
                paragraphs_map[key]["y0"] = min(paragraphs_map[key]["y0"], y)
                paragraphs_map[key]["x1"] = max(paragraphs_map[key]["x1"], x + w)
                paragraphs_map[key]["y1"] = max(paragraphs_map[key]["y1"], y + h)
            paragraphs_map[key]["words"].append(word)

        if not paragraphs_map:
            return []

        # Sort by vertical position — Tesseract already handles column order
        # within psm=3, so top-y sort gives correct reading order
        sorted_paras = sorted(paragraphs_map.values(), key=lambda p: p["top"])

        result = []
        for para in sorted_paras:
            text = " ".join(para["words"]).strip()
            text = re.sub(r" {2,}", " ", text)
            # Skip isolated page numbers / noise
            if re.match(r"^[-–—\s]*\d{1,4}[-–—\s]*$", text):
                continue
            if len(text) < 3:
                continue

            # Scale pixel bbox → PDF coordinate space (72 DPI)
            # img was rendered at 300 DPI; pdf_width/pdf_height are in PDF points
            if pdf_width and pdf_height and img_w and img_h:
                sx = pdf_width / img_w
                sy = pdf_height / img_h
                bbox = [
                    para["x0"] * sx,
                    para["y0"] * sy,
                    para["x1"] * sx,
                    para["y1"] * sy,
                ]
            else:
                bbox = [para["x0"], para["y0"], para["x1"], para["y1"]]

            result.append({
                "text": text,
                "bbox": bbox,
                "is_heading": _looks_like_heading(text),
                "word_count": len(para["words"]),
                "ocr": True,
            })
        return result

    except Exception:
        return []


def _ocr_simple(img: Image.Image) -> list:
    """
    Simple fallback: full-page OCR string split on blank lines.
    Used when image_to_data fails (e.g. Tesseract version too old).
    """
    try:
        raw = pytesseract.image_to_string(img, config="--psm 3 --oem 3").strip()
        if not raw:
            return []
        raw_paras = [p.strip() for p in re.split(r"\n{2,}", raw) if p.strip()]
        result = []
        for p in raw_paras:
            # Collapse single newlines within a paragraph into spaces
            text = re.sub(r"\n", " ", p)
            text = re.sub(r" {2,}", " ", text).strip()
            if re.match(r"^[-–—\s]*\d{1,4}[-–—\s]*$", text):
                continue
            if len(text) < 3:
                continue
            result.append({
                "text": text,
                "bbox": None,
                "is_heading": _looks_like_heading(text),
                "word_count": len(text.split()),
                "ocr": True,
            })
        return result
    except Exception:
        return []


def _looks_like_heading(text: str) -> bool:
    """Heuristic: short ALL-CAPS or Title Case lines are likely headings."""
    words = [w for w in text.split() if w.isalpha()]
    if not words:
        return False
    if len(words) > 12:
        return False
    if text.isupper():
        return True
    if all(w[0].isupper() for w in words):
        return True
    return False


def _extract_page_images(doc, page) -> list:
    """Extract embedded images from a PDF page via PyMuPDF."""
    images = []
    blocks = page.get_text("dict")["blocks"]
    for block in blocks:
        if block["type"] == 1:  # image block
            try:
                xref = block.get("image")
                if xref:
                    img_data = doc.extract_image(xref)
                    bbox = block.get("bbox", [0, 0, 0, 0])
                    images.append({
                        "bbox": bbox,
                        "index": len(images),
                        "format": img_data.get("ext", "png"),
                        "data_b64": base64.b64encode(img_data["image"]).decode(),
                    })
            except Exception:
                pass
    return images


def _extract_pdf_toc(doc) -> list:
    return [
        {"title": item[1], "page": item[2], "level": item[0]}
        for item in doc.get_toc()
    ]


# ── EPUB ──────────────────────────────────────────────────────────────────────

def _process_epub(file_bytes: bytes) -> dict:
    try:
        import ebooklib
        from ebooklib import epub
        from html.parser import HTMLParser
    except ImportError:
        raise RuntimeError("ebooklib not installed")

    book = epub.read_epub(io.BytesIO(file_bytes))
    pages = []
    total_words = 0
    toc = []

    class StructuredExtractor(HTMLParser):
        """
        Extract text preserving paragraph boundaries from HTML.
        Block-level tags (p, div, h1-h6, li, blockquote) become separate paragraphs.
        """
        BLOCK_TAGS = {"p", "div", "h1", "h2", "h3", "h4", "h5", "h6",
                      "li", "blockquote", "section", "article", "td", "th"}
        HEADING_TAGS = {"h1", "h2", "h3", "h4", "h5", "h6"}

        def __init__(self):
            super().__init__()
            self.paragraphs: list[dict] = []
            self._current_text: list[str] = []
            self._current_is_heading = False
            self._tag_stack: list[str] = []

        def handle_starttag(self, tag, attrs):
            tag = tag.lower()
            self._tag_stack.append(tag)
            if tag in self.BLOCK_TAGS:
                self._flush()
                self._current_is_heading = tag in self.HEADING_TAGS

        def handle_endtag(self, tag):
            tag = tag.lower()
            if tag in self.BLOCK_TAGS:
                self._flush()
            if self._tag_stack and self._tag_stack[-1] == tag:
                self._tag_stack.pop()

        def handle_data(self, data):
            text = data.strip()
            if text:
                self._current_text.append(text)

        def _flush(self):
            text = " ".join(self._current_text).strip()
            text = re.sub(r" {2,}", " ", text)
            if text and len(text) > 2:
                self.paragraphs.append({
                    "text": text,
                    "bbox": None,
                    "is_heading": self._current_is_heading,
                    "word_count": len(text.split()),
                })
            self._current_text = []
            self._current_is_heading = False

        def get_paragraphs(self):
            self._flush()
            return self.paragraphs

    for i, item in enumerate(book.get_items_of_type(ebooklib.ITEM_DOCUMENT)):
        extractor = StructuredExtractor()
        extractor.feed(item.get_body_content().decode("utf-8", errors="ignore"))
        paragraphs = extractor.get_paragraphs()

        if not paragraphs:
            continue

        total_words += sum(p["word_count"] for p in paragraphs)

        # Use first heading as TOC entry if available
        heading = next((p["text"] for p in paragraphs if p["is_heading"]), item.get_name())
        toc.append({"title": heading, "page": i + 1, "level": 1})
        pages.append({"page": i + 1, "paragraphs": paragraphs, "images": []})

    return {"pages": pages, "total_pages": len(pages), "total_words": total_words, "toc": toc}


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _process_docx(file_bytes: bytes) -> dict:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx not installed")

    doc = Document(io.BytesIO(file_bytes))
    pages = [{"page": 1, "paragraphs": [], "images": []}]
    total_words = 0
    chars_per_page = 3000
    char_count = 0
    toc = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        is_heading = para.style.name.startswith("Heading")
        if is_heading:
            toc.append({"title": text, "page": len(pages), "level": 1})
        word_count = len(text.split())
        total_words += word_count
        pages[-1]["paragraphs"].append({
            "text": text,
            "bbox": None,
            "is_heading": is_heading,
            "word_count": word_count,
        })
        char_count += len(text)
        if char_count > chars_per_page:
            pages.append({"page": len(pages) + 1, "paragraphs": [], "images": []})
            char_count = 0

    return {"pages": pages, "total_pages": len(pages), "total_words": total_words, "toc": toc}


# ── TXT ───────────────────────────────────────────────────────────────────────

def _process_txt(file_bytes: bytes) -> dict:
    text = file_bytes.decode("utf-8", errors="ignore")
    raw_paras = [p.strip() for p in text.split("\n\n") if p.strip()]
    chars_per_page = 3000
    pages = []
    current_page_paras = []
    current_chars = 0
    total_words = 0

    for para_text in raw_paras:
        wc = len(para_text.split())
        total_words += wc
        current_page_paras.append({"text": para_text, "bbox": None, "is_heading": False, "word_count": wc})
        current_chars += len(para_text)
        if current_chars > chars_per_page:
            pages.append({"page": len(pages) + 1, "paragraphs": current_page_paras, "images": []})
            current_page_paras = []
            current_chars = 0

    if current_page_paras:
        pages.append({"page": len(pages) + 1, "paragraphs": current_page_paras, "images": []})

    return {"pages": pages, "total_pages": len(pages), "total_words": total_words, "toc": []}


# ── Celery task ───────────────────────────────────────────────────────────────

@celery_app.task(name="process_book", bind=True, max_retries=3)
def process_book_task(self, book_id: str, s3_key: str, file_type: str):
    """Async Celery task: download from storage, parse, index, update DB."""
    import asyncio
    from core.database import AsyncSessionLocal
    from models.db import Book
    from services.storage_service import StorageService
    from services.rag_service import index_book
    from sqlalchemy import select, update

    async def _run():
        storage = StorageService()
        file_bytes = await storage.download(s3_key, book_id=book_id)
        result = process_document(file_bytes, file_type)

        # Index for RAG
        try:
            index_book(book_id, result["pages"])
        except Exception:
            pass  # RAG indexing is best-effort

        async with AsyncSessionLocal() as db:
            await db.execute(
                update(Book).where(Book.id == book_id).values(
                    extracted_text=result,
                    total_pages=result["total_pages"],
                    total_words=result["total_words"],
                    toc=result["toc"],
                    vector_index_id=book_id,
                    status="ready",
                )
            )
            await db.commit()

    try:
        asyncio.run(_run())
    except Exception as exc:
        import asyncio as _asyncio
        from core.database import AsyncSessionLocal
        from models.db import Book
        from sqlalchemy import update as _update

        async def _mark_error():
            async with AsyncSessionLocal() as db:
                await db.execute(
                    _update(Book).where(Book.id == book_id).values(
                        status="error", error_message=str(exc)
                    )
                )
                await db.commit()

        _asyncio.run(_mark_error())
        raise self.retry(exc=exc, countdown=30)
